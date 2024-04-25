# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import random
import string

sumar_txt = 'Сегодня мы будем говорить о массивах. На прошлой лекции мы познакомились с языком программирования C-Sharp, научились решать базовые задачи и попрактиковались в разработке циклических условных конструкций, а также арифметических операций. (00:00:0.00 - 00:00:27.68) Сегодня мы вспомним основные понятия, связанные с массивами, их характеристики и возможные области применения. Массив представляет собой структуру данных для хранения элементов одного типа. Массивы могут состоять из целых чисел, вещественных, символов или строк, и их можно сравнить с упорядоченным набором инструментов или набором елочных игрушек, обладающими характеристикой упорядоченности. (00:00:45.24 - 00:00:56.96) Важным элементом массива является индекс элемента массива, который позволяет обращаться к отдельным элементам. Мы также рассмотрим операции создания, заполнения и вывода массивов на экран. Ближе к концу лекции мы познакомимся с циклами For и For Each, а также обсудим значимость изучения английского языка для программистов. (00:00:56.96 - 00:01:40.60)'


def txt_spliter(txt):
    pattern = r'\(\d+:\d+:\d+\.\d+ - \d+:\d+:\d+\.\d+\)'
    blocks = re.split(pattern, txt)

    blocks = [block.strip() for block in blocks if block.strip()]

    splited_txt = []

    for i, block in enumerate(blocks, start=1):
        splited_txt.append((f"Блок {i}", block))

    return splited_txt

def segmentation_for_sumar(sumar_txt):
    result = txt_spliter(sumar_txt)

    pages = []
    page_titles = []
    for segments in result:
        page_titles.append(segments[0])
        pages.append(segments[1])
    return [pages,page_titles]

def create_word_document(pages, page_titles):
    doc = Document()

    for page_num, (page_content, page_title) in enumerate(zip(pages, page_titles), start=1):
        # Добавляем новую страницу
        if page_num > 1:
            doc.add_page_break()

        # Добавляем заголовок страницы
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(page_title)
        title_run.bold = True
        title_run.font.size = Pt(14)  # Установите желаемый размер
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Выравнивание по центру

        # Добавляем текст на страницу
        doc.add_paragraph(page_content)
    name=generate_random_name()
    # Сохраняем документ
    doc.save(name)
    return name

def generate_random_name():
    random_name = ''.join(random.choices(string.ascii_letters + string.digits, k=20))
    return random_name + ".docx"
def txt2docx(sumar_txt):
    pages=segmentation_for_sumar(sumar_txt)

    name = create_word_document(pages[0], pages[1])
    return name

#print(txt2docx(sumar_txt))